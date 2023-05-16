import tkinter as tk
from tkinter import ttk, messagebox, filedialog, scrolledtext
import docx
from PIL import ImageGrab, ImageTk, Image
from io import BytesIO
import clipboard

class ImageLabel(tk.Label):
    def __init__(self, master=None, **kwargs):
        super().__init__(master, **kwargs)
        self.bind("<B1-Motion>", self.drag)
        self.bind("<Button-1>", self.click)
        self.bind("<ButtonRelease-1>", self.drop)

    def click(self, event):
        self.startX = event.x
        self.startY = event.y

    def drag(self, event):
        x = self.winfo_x() - self.startX + event.x
        y = self.winfo_y() - self.startY + event.y
        self.place(x=x, y=y)

    def drop(self, event):
        x = self.winfo_x() + event.x
        y = self.winfo_y() + event.y
        self.place(x=x, y=y)

        # This is where you would handle dropping the image label
        
        pass

def paste_screenshot(self, event):
    try:
        image = ImageGrab.grabclipboard()
        if isinstance(image, Image.Image):  # Make sure it's an image
            # Save the image and add it to the list
            filename = f"image_{len(self.images)}.png"
            image.save(filename)
            self.images.append(filename)

            # Display the image in the application
            image.thumbnail((100, 100))  # Reduce the size of the image
            photo = ImageTk.PhotoImage(image)
            img_label = ImageLabel(self.content_frame, image=photo)  # Use ImageLabel instead of tk.Label
            img_label.image = photo
            img_label.place(x=100, y=100)  # Use place instead of grid to allow moving the label
        else:
            print("No image in the clipboard.")
    except Exception as e:
        print(e)  # This will print the actual error message
        messagebox.showerror('Error', 'Could not paste the image!')

class ScrollableFrame(tk.Frame):
    def __init__(self, master=None, **kwargs):
        super().__init__(master, **kwargs)

        # Create a canvas object and a vertical scrollbar for scrolling it
        self.vscrollbar = tk.Scrollbar(self, orient=tk.VERTICAL)
        self.vscrollbar.pack(fill=tk.Y, side=tk.RIGHT, expand=False)
        self.canvas = tk.Canvas(self, bd=0, highlightthickness=0, yscrollcommand=self.vscrollbar.set)
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.vscrollbar.config(command=self.canvas.yview)

        # Reset the view
        self.canvas.xview_moveto(0)
        self.canvas.yview_moveto(0)

        # Create a frame inside the canvas which will be scrolled with it
        self.interior = tk.Frame(self.canvas)
        self.canvas.create_window((0, 0), window=self.interior, anchor=tk.NW)

        self.interior.bind('<Configure>', self.configure_interior)
        self.canvas.bind('<Configure>', self.configure_canvas)

    def configure_interior(self, event):
        # Update the scrollbars to match the size of the inner frame
        size = (self.interior.winfo_reqwidth(), self.interior.winfo_reqheight())
        self.canvas.config(scrollregion="0 0 %s %s" % size)
        if self.interior.winfo_reqwidth() != self.canvas.winfo_width():
            # Update the canvas's width to fit the inner frame
            self.canvas.config(width=self.interior.winfo_reqwidth())

    def configure_canvas(self, event):
        if self.interior.winfo_reqwidth() != self.canvas.winfo_width():
            # Update the inner frame's width to fill the canvas
            self.canvas.itemconfigure('interior', width=self.canvas.winfo_width())

class IssueReportingApp(tk.Tk):
    def __init__(self):
        super().__init__()

        # Initialize UI elements
        self.init_ui()
        self.image_frame = ScrollableFrame(self)
        self.image_frame.pack(side='right', fill='both', expand=True)


        # List to store images
        self.images = []

    def init_ui(self):
        frame = tk.Frame(self)
        frame.pack(fill='both', expand=True)

        canvas = tk.Canvas(frame)
        canvas.pack(side='left', fill='both', expand=True)

        scrollbar = ttk.Scrollbar(frame, orient='vertical', command=canvas.yview)
        scrollbar.pack(side='right', fill='y')

        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.bind('<Configure>', lambda e: canvas.configure(scrollregion=canvas.bbox('all')))

        self.content_frame = tk.Frame(canvas)
        canvas.create_window((0, 0), window=self.content_frame, anchor='nw')

        # Create input fields for the required information
        ttk.Label(self.content_frame, text="Req-Number").grid(row=0, column=0, sticky='w', pady=5)
        self.req_number = ttk.Entry(self.content_frame)
        self.req_number.grid(row=0, column=1, pady=5)

        # (Continue with rest of your code, replacing content_frame with self.content_frame)
        ttk.Label(self.content_frame, text="Company").grid(row=1, column=0, sticky='w', pady=5)
        self.company = ttk.Entry(self.content_frame)
        self.company.grid(row=1, column=1, pady=5)

        ttk.Label(self.content_frame, text="Country").grid(row=2, column=0, sticky='w', pady=5)
        self.country = ttk.Combobox(self.content_frame, values=['USA', 'UK', 'Germany', 'France'])  # Add more countries as needed
        self.country.grid(row=2, column=1, pady=5)

        ttk.Label(self.content_frame, text="Account name").grid(row=3, column=0, sticky='w', pady=5)
        self.account_name = ttk.Entry(self.content_frame)
        self.account_name.grid(row=3, column=1, pady=5)

        ttk.Label(self.content_frame, text="Impacted user's ID(s)").grid(row=4, column=0, sticky='w', pady=5)
        self.user_ids = ttk.Entry(self.content_frame)
        self.user_ids.grid(row=4, column=1, pady=5)

        ttk.Label(self.content_frame, text="Environment, Application, Sub Product, Dataset").grid(row=5, column=0, sticky='w', pady=5)
        self.dataset = ttk.Entry(self.content_frame)
        self.dataset.grid(row=5, column=1, pady=5)

        ttk.Label(self.content_frame, text="Report Details").grid(row=6, column=0, sticky='w', pady=5)
        self.report_details = scrolledtext.ScrolledText(self.content_frame, width=40, height=5)
        self.report_details.grid(row=6, column=1, pady=5)

        ttk.Label(self.content_frame, text="Is the issue replicable?").grid(row=7, column=0, sticky='w', pady=5)
        self.replicable = ttk.Combobox(self.content_frame, values=['YES', 'NO'])
        self.replicable.grid(row=7, column=1, pady=5)

        ttk.Label(self.content_frame, text="Steps/Troubleshooting").grid(row=8, column=0, sticky='w', pady=5)
        self.steps = scrolledtext.ScrolledText(self.content_frame, width=40, height=5)
        self.steps.grid(row=8, column=1, pady=5)

        ttk.Label(self.content_frame, text="Time and timezone of error").grid(row=9, column=0, sticky='w', pady=5)
        self.error_time = ttk.Entry(self.content_frame)
        self.error_time.grid(row=9, column=1, pady=5)

        ttk.Label(self.content_frame, text="Describe the issue").grid(row=10, column=0, sticky='w', pady=5)
        self.issue_description = scrolledtext.ScrolledText(self.content_frame, width=40, height=5)
        self.issue_description.grid(row=10, column=1, pady=5)

        # Image display area
        self.image_label = tk.Label(self.content_frame)  
        self.image_label.grid(row=11, column=0, columnspan=2)

    # Bind Ctrl+V for pasting images
        self.bind('<Control-v>', self.paste_screenshot)

        # Create buttons for generating output
        self.generate_button = ttk.Button(self.content_frame, text="Generate Word Document", command=self.generate_word_document)
        self.generate_button.grid(row=12, column=0, columnspan=2, pady=5)

        self.copy_output_button = ttk.Button(self.content_frame, text="Copy Output", command=self.generate_copy_ready_text)
        self.copy_output_button.grid(row=13, column=0, columnspan=2, pady=5)

        # Create a button for exiting the application
        self.exit_button = ttk.Button(self.content_frame, text="Exit", command=self.exit_application)
        self.exit_button.grid(row=14, column=0, columnspan=2, pady=5)

    def generate_word_document(self):
        doc = docx.Document()

        # Add content to the document
        doc.add_paragraph(f"Req-Number: {self.req_number.get()}")
        doc.add_paragraph(f"Company: {self.company.get()}")
        doc.add_paragraph(f"Country: {self.country.get()}")
        # ... (Add other input fields)

        # Add images to the document
        for image in self.images:
            doc.add_picture(image)

        # Save the Word document
        try:
            doc.save(self.req_number.get() + '.docx')
            messagebox.showinfo('Info', 'Word document created successfully!')
        except Exception as e:
            messagebox.showerror('Error', 'Could not save the Word document!')

    def generate_copy_ready_text(self):
        # Prepare text using input fields
        text = f"Req-Number: {self.req_number.get()}\n"
        text += f"Company: {self.company.get()}\n"
        text += f"Country: {self.country.get()}\n"
        # ... (Add other input fields)

        # Copy text to clipboard
        clipboard.copy(text)
        messagebox.showinfo('Info', 'Text copied to clipboard!')

def paste_screenshot(self, event):
    try:
        image = ImageGrab.grabclipboard()
        if isinstance(image, Image.Image):  # Make sure it's an image
            # Save the image and add it to the list
            filename = f"image_{len(self.images)}.png"
            image.save(filename)
            self.images.append(filename)

            # Display the image in the application
            image.thumbnail((100, 100))  # Reduce the size of the image
            photo = ImageTk.PhotoImage(image)
            img_label = ImageLabel(self.content_frame, image=photo)  # Use ImageLabel instead of tk.Label
            img_label.image = photo
            img_label.place(x=100, y=100)  # Use place instead of grid to allow moving the label
        else:
            print("No image in the clipboard.")
    except Exception as e:
        print(e)  # This will print the actual error message
        messagebox.showerror('Error', 'Could not paste the image!')


    def exit_application(self):
        msg_box = messagebox.askquestion('Exit Application', 'Are you sure you want to exit the application?', icon='warning')
        if msg_box == 'yes':
            self.destroy()
        else:
            messagebox.showinfo('Return', 'You will now return to the application screen')

    

if __name__ == '__main__':
    app = IssueReportingApp()
    app.mainloop()


