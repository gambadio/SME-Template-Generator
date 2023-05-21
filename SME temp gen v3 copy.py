import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext
import docx
from PIL import ImageGrab, ImageTk, Image
from io import BytesIO
import clipboard
import os
import re
import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext, filedialog
import shutil 
from docx.shared import Inches
import win32com.client

class CustomScrolledText(scrolledtext.ScrolledText):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.bind('<Control-Shift-v>', self.paste_with_braces)

    def paste_with_braces(self, event):
        clipboard_content = clipboard.paste()
        if os.path.isfile(clipboard_content):
            self.insert(tk.INSERT, f"{{{clipboard_content}}}\n")
        else:
            self.insert(tk.INSERT, clipboard_content)



    def paste(self, event):
        clipboard_content = clipboard.paste()
        if os.path.isfile(clipboard_content):
            self.insert(tk.INSERT, f"{{{clipboard_content}}}\n")
        else:
            self.insert(tk.INSERT, clipboard_content)


class IssueReportingApp(tk.Tk):

    

    def on_content_frame_configure(self, event):
        # Set scroll region of Canvas widget
        self.canvas.configure(scrollregion=self.canvas.bbox('all'))

    def on_canvas_configure(self, event):
        # Resize inner window of Canvas widget
        self.canvas.itemconfig(1, width=event.width)

    def on_mouse_wheel(self, event):
        # Scroll Canvas widget vertically using mouse wheel
        self.canvas.yview_scroll(-int(event.delta/120), 'units')


    def __init__(self):

        # List to store images
        self.images = []

        super().__init__()

        # Set window title
        self.title("Issue Reporting")

        # Create a Canvas widget and a Scrollbar widget
        self.canvas = tk.Canvas(self)
        self.scrollbar = ttk.Scrollbar(self, orient='vertical', command=self.canvas.yview)
        self.canvas.configure(yscrollcommand=self.scrollbar.set)

        # Create a new frame to hold all the content of the app
        self.content_frame = tk.Frame(self.canvas)

        # Create a window inside the Canvas to hold the content_frame
        self.canvas.create_window((0, 0), window=self.content_frame, anchor='nw')

        # Pack the Canvas and Scrollbar widgets
        self.canvas.pack(side='left', fill='both', expand=True)
        self.scrollbar.pack(side='right', fill='y')

        # Initialize UI elements
        self.init_ui()

        # Bind events for scrolling
        self.content_frame.bind('<Configure>', self.on_content_frame_configure)
        self.canvas.bind('<Configure>', self.on_canvas_configure)
        self.canvas.bind_all('<MouseWheel>', self.on_mouse_wheel)

    def init_ui(self):

        # Set window title
        self.title("SME Template")

        # Set initial size of main window
        self.geometry('800x700')

        # Create input fields for the required information
        ttk.Label(self.content_frame, text="Req-Number").grid(row=0, column=0, sticky='w', pady=5)
        self.req_number = ttk.Entry(self.content_frame)
        # Add padding to the left side of the req_number field
        self.req_number.grid(row=0, column=1, pady=5, padx=(100,0))

        ttk.Label(self.content_frame, text="Country").grid(row=2, column=0, sticky='w', pady=5)
        self.country = ttk.Entry(self.content_frame) 
        # Add padding to the left side of the country field
        self.country.grid(row=2, column=1, pady=5, padx=(100,0))



        ttk.Label(self.content_frame, text="Account name").grid(row=3, column=0, sticky='w', pady=5)
        self.account_name = ttk.Entry(self.content_frame)
        self.account_name.grid(row=3, column=1, pady=5, padx=(100,0))

        ttk.Label(self.content_frame, text="Impacted user's ID(s)").grid(row=4, column=0, sticky='w', pady=5)
        self.user_ids = ttk.Entry(self.content_frame)
        self.user_ids.grid(row=4, column=1, pady=5, padx=(100,0))

        # Split label into two lines
        ttk.Label(self.content_frame, text="Environment, Application,\nSub Product,\nDataset").grid(row=5,column=0,pady=(10,0),sticky='w')
        # Make dataset field an Entry field with a larger width
        self.dataset = CustomScrolledText(self.content_frame, width=40, height=2)
        self.dataset.grid(row = 5,column = 1,pady= 5, padx=(100,0))

        ttk.Label(self.content_frame, text="Report Details").grid(row=6, column=0, sticky='w', pady=5)
        self.report_details = CustomScrolledText(self.content_frame, width=40, height=5)
        self.report_details.grid(row=6, column=1, pady=5, padx=(100,0))

        ttk.Label(self.content_frame, text="Is the issue replicable?").grid(row=7, column=0, sticky='w', pady=5)
        self.replicable = ttk.Combobox(self.content_frame, values=['YES', 'NO'])
        self.replicable.grid(row=7, column=1, pady=5, padx=(100,0))

        ttk.Label(self.content_frame, text="Steps/Troubleshooting").grid(row=8, column=0, sticky='w', pady=5)
        self.steps = CustomScrolledText(self.content_frame, width=40, height=5)
        self.steps.grid(row=8, column=1, pady=5, padx=(100,0))

        ttk.Label(self.content_frame, text="Time and timezone of error").grid(row=9, column=0, sticky='w', pady=5)
        self.error_time = ttk.Entry(self.content_frame)
        self.error_time.grid(row=9,column=1, pady=5, padx=(100,0))

        ttk.Label(self.content_frame, text="Describe the issue").grid(row=10, column=0, sticky='w', pady=5)
        self.issue_description = CustomScrolledText(self.content_frame, width=40, height=5)
        self.issue_description.grid(row=10, column=1, pady=5, padx=(100,0))

        # Image display area
        self.image_label = tk.Label(self.content_frame) 
        self.image_label.grid(row=11, column=0, columnspan=2)

        # Bind Ctrl+V for pasting images
        self.bind('<Control-v>', self.paste_screenshot)

        # Create buttons for generating output
        self.generate_button = ttk.Button(self.content_frame, text="Generate Word Document", command=self.generate_word_document)
        self.generate_button.grid(row=12, column=0, columnspan=2, pady=5)

        self.copy_output_button = ttk.Button(self.content_frame, text="Copy Output", command=self.generate_copy_ready_text)
        self.copy_output_button.grid(row=13, column=0, columnspan=2, pady=5)

        # Create a button for exiting the application
        self.exit_button = ttk.Button(self.content_frame, text="Exit", command=self.exit_application)
        self.exit_button.grid(row=14,column=0, columnspan=2, pady=5)

    def paste_screenshot(self, event):
        try:
            clipboard_content = clipboard.paste()
            # Replace forward slashes with backslashes in the file path
            clipboard_content = clipboard_content.replace('/', '\\')
            if os.path.isfile(clipboard_content):
                # Open the image file and display it
                image = Image.open(clipboard_content)
                # Save the image and add it to the list
                new_filename = os.path.abspath(f"image_{len(self.images)}.png")
                image.save(new_filename)
                self.images.append(new_filename)

                # Copy the image file to the application's folder
                shutil.copy(clipboard_content, new_filename)

                # Insert placeholder into active text field
                active_widget = self.focus_get()
                if isinstance(active_widget, CustomScrolledText):
                    active_widget.insert(tk.END, f"\n{{{new_filename}}}\n") # Always wrap filename in curly braces
            else:
                # Handle pasted image data
                image = ImageGrab.grabclipboard()
                if isinstance(image, Image.Image): # Make sure it's an image
                    # Save the image and add it to the list
                    filename = os.path.abspath(f"image_{len(self.images)}.png")
                    image.save(filename)
                    self.images.append(filename)

                    # Insert placeholder into active text field
                    active_widget = self.focus_get()
                    if isinstance(active_widget, CustomScrolledText):
                        active_widget.insert(tk.END, f"\n{{{filename}}}\n") # Always wrap filename in curly braces
        except Exception as e:
            print(e) # This will print the actual error message
            messagebox.showerror('Error', 'Could not paste the image!')


    def add_image(self):
        filename = filedialog.askopenfilename(initialdir="/", title="Select Image", filetypes=(("jpeg files", "*.jpg"), ("png files", "*.png")))
        if filename:
            image = Image.open(filename)
            image.thumbnail((200, 200))
            photo = ImageTk.PhotoImage(image)

            self.images.append(filename)


    def generate_word_document(self):
        doc = docx.Document()

        # Add content to the document
        doc.add_paragraph(f"Req-Number: {self.req_number.get()}")
        doc.add_paragraph(f"Country: {self.country.get()}")
        doc.add_paragraph(f"Account name: {self.account_name.get()}")
        doc.add_paragraph(f"Impacted user's ID(s): {self.user_ids.get()}")

        # Add "Report Details:" before parsing text from report_details attribute
        report_details_para = doc.add_paragraph("Environment, Application, Sub Product, Dataset: ")
        # Parse text from report_details text field and insert images at appropriate locations
        text = self.dataset.get('1.0', 'end')
        lines = text.split('\n')
        for line in lines:
            # Check if line is a file path wrapped in curly braces
            if line.startswith('{') and line.endswith('}'):
                image_filename = line[1:-1]  # Remove the curly braces
            else:
                image_filename = line  # Use the line as is

            # Replace forward slashes with backslashes in the file path
            image_filename = image_filename.replace('/', '\\')

            # Create a new run in the report_details_para paragraph
            run = report_details_para.add_run()

            if os.path.isfile(image_filename):
                # Add the image to the run
                run.add_picture(image_filename, width=Inches(6.0))  # Adjust the width as needed
            else:
                # Add the text to the run
                run.add_text(line + '\n')

        # Add "Report Details:" before parsing text from report_details attribute
        report_details_para = doc.add_paragraph("Report Details: ")
        # Parse text from report_details text field and insert images at appropriate locations
        text = self.report_details.get('1.0', 'end')
        lines = text.split('\n')
        for line in lines:
            # Check if line is a file path wrapped in curly braces
            if line.startswith('{') and line.endswith('}'):
                image_filename = line[1:-1]  # Remove the curly braces
            else:
                image_filename = line  # Use the line as is

            # Replace forward slashes with backslashes in the file path
            image_filename = image_filename.replace('/', '\\')

            # Create a new run in the report_details_para paragraph
            run = report_details_para.add_run()

            if os.path.isfile(image_filename):
                # Add the image to the run
                run.add_picture(image_filename, width=Inches(6.0))  # Adjust the width as needed
            else:
                # Add the text to the run
                run.add_text(line + '\n')

        doc.add_paragraph(f"Is the issue replicable? {self.replicable.get()}")

        # Add "Steps/Troubleshooting:" before parsing text from steps attribute
        steps_para = doc.add_paragraph("Steps/Troubleshooting: ")
        # Parse text from steps text field and insert images at appropriate locations
        text = self.steps.get('1.0', 'end')
        lines = text.split('\n')
        for line in lines:
            # Check if line is a file path wrapped in curly braces
            if line.startswith('{') and line.endswith('}'):
                image_filename = line[1:-1]  # Remove the curly braces
            else:
                image_filename = line  # Use the line as is

            # Replace forward slashes with backslashes in the file path
            image_filename = image_filename.replace('/', '\\')

            # Create a new run in the steps_para paragraph
            run = steps_para.add_run()

            if os.path.isfile(image_filename):
                # Add the image to the run
                run.add_picture(image_filename, width=Inches(6.0))  # Adjust the width as needed
            else:
                # Add the text to the run
                run.add_text(line + '\n')

        doc.add_paragraph(f"Time and timezone of error: {self.error_time.get()}")

        # Add "Describe the issue:" before parsing text from issue_description attribute
        issue_description_para = doc.add_paragraph("Describe the issue: ")
        # Parse text from issue_description text field and insert images at appropriate locations
        text = self.issue_description.get('1.0', 'end')
        lines = text.split('\n')
        for line in lines:
            # Check if line is a file path wrapped in curly braces
            if line.startswith('{') and line.endswith('}'):
                image_filename = line[1:-1]  # Remove the curly braces
            else:
                image_filename = line  # Use the line as is

            # Replace forward slashes with backslashes in the file path
            image_filename = image_filename.replace('/', '\\')

            # Create a new run in the issue_description_para paragraph
            run = issue_description_para.add_run()

            if os.path.isfile(image_filename):
                # Add the image to the run
                run.add_picture(image_filename, width=Inches(6.0))  # Adjust the width as needed
            else:
                # Add the text to the run
                run.add_text(line + '\n')

        # Save the Word document
        try:
            doc.save(self.req_number.get() + '.docx')
            messagebox.showinfo('Info', 'Word document created successfully!')
        except Exception as e:
            messagebox.showerror('Error', 'Could not save the Word document!')


    def generate_copy_ready_text(self):
        doc = docx.Document()

        # Add content to the document
        doc.add_paragraph(f"Req-Number: {self.req_number.get()}")
        doc.add_paragraph(f"Country: {self.country.get()}")
        doc.add_paragraph(f"Account name: {self.account_name.get()}")
        doc.add_paragraph(f"Impacted user's ID(s): {self.user_ids.get()}")

        # Add "Report Details:" before parsing text from report_details attribute
        report_details_para = doc.add_paragraph("Environment, Application, Sub Product, Dataset: ")
        # Parse text from report_details text field and insert images at appropriate locations
        text = self.dataset.get('1.0', 'end')
        lines = text.split('\n')
        for line in lines:
            # Check if line is a file path wrapped in curly braces
            if line.startswith('{') and line.endswith('}'):
                image_filename = line[1:-1]  # Remove the curly braces
            else:
                image_filename = line  # Use the line as is

            # Replace forward slashes with backslashes in the file path
            image_filename = image_filename.replace('/', '\\')

            # Create a new run in the report_details_para paragraph
            run = report_details_para.add_run()

            if os.path.isfile(image_filename):
                # Add the image to the run
                run.add_picture(image_filename, width=Inches(6.0))  # Adjust the width as needed
            else:
                # Add the text to the run
                run.add_text(line + '\n')

        # Add "Report Details:" before parsing text from report_details attribute
        report_details_para = doc.add_paragraph("Report Details: ")
        # Parse text from report_details text field and insert images at appropriate locations
        text = self.report_details.get('1.0', 'end')
        lines = text.split('\n')
        for line in lines:
            # Check if line is a file path wrapped in curly braces
            if line.startswith('{') and line.endswith('}'):
                image_filename = line[1:-1]  # Remove the curly braces
            else:
                image_filename = line  # Use the line as is

            # Replace forward slashes with backslashes in the file path
            image_filename = image_filename.replace('/', '\\')

            # Create a new run in the report_details_para paragraph
            run = report_details_para.add_run()

            if os.path.isfile(image_filename):
                # Add the image to the run
                run.add_picture(image_filename, width=Inches(6.0))  # Adjust the width as needed
            else:
                # Add the text to the run
                run.add_text(line + '\n')

        doc.add_paragraph(f"Is the issue replicable? {self.replicable.get()}")

        # Add "Steps/Troubleshooting:" before parsing text from steps attribute
        steps_para = doc.add_paragraph("Steps/Troubleshooting: ")
        # Parse text from steps text field and insert images at appropriate locations
        text = self.steps.get('1.0', 'end')
        lines = text.split('\n')
        for line in lines:
            # Check if line is a file path wrapped in curly braces
            if line.startswith('{') and line.endswith('}'):
                image_filename = line[1:-1]  # Remove the curly braces
            else:
                image_filename = line  # Use the line as is

            # Replace forward slashes with backslashes in the file path
            image_filename = image_filename.replace('/', '\\')

            # Create a new run in the steps_para paragraph
            run = steps_para.add_run()

            if os.path.isfile(image_filename):
                # Add the image to the run
                run.add_picture(image_filename, width=Inches(6.0))  # Adjust the width as needed
            else:
                # Add the text to the run
                run.add_text(line + '\n')

        doc.add_paragraph(f"Time and timezone of error: {self.error_time.get()}")

        # Add "Describe the issue:" before parsing text from issue_description attribute
        issue_description_para = doc.add_paragraph("Describe the issue: ")
        # Parse text from issue_description text field and insert images at appropriate locations
        text = self.issue_description.get('1.0', 'end')
        lines = text.split('\n')
        for line in lines:
            # Check if line is a file path wrapped in curly braces
            if line.startswith('{') and line.endswith('}'):
                image_filename = line[1:-1]  # Remove the curly braces
            else:
                image_filename = line  # Use the line as is

            # Replace forward slashes with backslashes in the file path
            image_filename = image_filename.replace('/', '\\')

            # Create a new run in the issue_description_para paragraph
            run = issue_description_para.add_run()

            if os.path.isfile(image_filename):
                # Add the image to the run
                run.add_picture(image_filename, width=Inches(6.0))  # Adjust the width as needed
            else:
                # Add the text to the run
                run.add_text(line + '\n')

        # Save the Word document to a temporary file
        temp_filename = 'temp.docx'
        doc.save(temp_filename)

        # Open the Word document and copy its content to the clipboard
        word = win32com.client.Dispatch('Word.Application')
        doc = word.Documents.Open(os.path.abspath(temp_filename))
        doc.Content.Copy()
        doc.Close()

        # Delete the temporary file
        os.remove(temp_filename)

        messagebox.showinfo('Info', 'Content copied to clipboard!')

    def exit_application(self):
        msg_box = messagebox.askquestion('Exit Application', 'Are you sure you want to exit the application?', icon='warning')
        if msg_box == 'yes':
            self.destroy()
        else:
            messagebox.showinfo('Return', 'You will now return to the application screen')


if __name__ == '__main__':
    app = IssueReportingApp()
    app.mainloop()


