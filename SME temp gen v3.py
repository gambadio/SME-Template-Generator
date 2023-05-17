import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext
import docx
from PIL import ImageGrab, ImageTk, Image
from io import BytesIO
import clipboard
import os
import re
import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext, filedialog
import shutil 
from docx.shared import Inches
import win32com.client

class CustomScrolledText(scrolledtext.ScrolledText):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.bind('<Control-Shift-v>', self.paste_with_braces)

    def paste_with_braces(self, event):
        clipboard_content = clipboard.paste()
        if os.path.isfile(clipboard_content):
            self.insert(tk.INSERT, f"{{{clipboard_content}}}\n")
        else:
            self.insert(tk.INSERT, clipboard_content)



    def paste(self, event):
        clipboard_content = clipboard.paste()
        if os.path.isfile(clipboard_content):
            self.insert(tk.INSERT, f"{{{clipboard_content}}}\n")
        else:
            self.insert(tk.INSERT, clipboard_content)


class ImageLabel(tk.Label):
    def __init__(self, master=None, image_filename=None, app=None, **kwargs):
        super().__init__(master, **kwargs)
        self.image_filename = image_filename
        self.app = app  # Store the app instance
        self.bind('<Button-3>', self.show_context_menu)

        self.context_menu = tk.Menu(self, tearoff=0)
        self.context_menu.add_command(label="Delete", command=self.delete_image)


    def show_context_menu(self, event):
        self.context_menu.tk_popup(event.x_root, event.y_root)

    def delete_image(self):
        # Get the IssueReportingApp instance
        app = self._root()

        # Remove the image from the list of images in the IssueReportingApp instance
        if self.image_filename in app.images:
            app.images.remove(self.image_filename)

        # Delete the image file
        os.remove(self.image_filename)

        # Delete all instances of the placeholder in all text fields
        for widget in app.content_frame.winfo_children():
            if isinstance(widget, (CustomScrolledText, DragDropText)):
                text = widget.get('1.0', 'end')
                text = re.sub(f"{{{self.image_filename}}}\n", "", text)  # Use regular expression to replace all instances of placeholder with empty string
                widget.delete('1.0', 'end')
                widget.insert('1.0', text)

        # Delete corresponding ImageLabel in preview
        for widget in app.image_frame.scrollable_frame.winfo_children():
            if isinstance(widget, ImageLabel) and widget.image_filename == self.image_filename:
                widget.destroy()

        # Destroy the ImageLabel widget
        self.destroy()




class DragDropText(CustomScrolledText):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)

        self.drag_start_pos = None
        self.dragged_text = None

        self.bind('<ButtonPress-1>', self.drag_start)
        self.bind('<B1-Motion>', self.drag)
        self.bind('<ButtonRelease-1>', self.drop)

    def drag_start(self, event):
        # Get the index of the character under the cursor
        index = self.index(f'@{event.x},{event.y}')
        # Check if the character is part of a placeholder
        line = self.get(f'{index} linestart', f'{index} lineend')
        start_index = line.find('{')
        end_index = line.find('}')
        if start_index != -1 and end_index != -1:
            # Select the placeholder
            self.tag_add('sel', f'{index} linestart+{start_index}c', f'{index} linestart+{end_index+1}c')
            # Store the start position and dragged text
            self.drag_start_pos = f'{index} linestart+{start_index}c'
            self.dragged_text = self.selection_get()

    def drag(self, event):
        if self.drag_start_pos:
            # Remove previous insertion mark
            self.mark_unset('insert')
            # Set new insertion mark at cursor position
            index = self.index(f'@{event.x},{event.y}')
            self.mark_set('insert', index)

    def drop(self, event):
        if self.drag_start_pos:
            # Get the drop position
            index = self.index(f'@{event.x},{event.y}')
            # Delete the selected text
            self.delete(self.drag_start_pos, f'{self.drag_start_pos}+{len(self.dragged_text)}c')
            # Insert the dragged text at the drop position
            self.insert(index, self.dragged_text)
            # Clear selection and insertion mark
            self.tag_remove('sel', '1.0', 'end')
            self.mark_unset('insert')
            # Reset drag state
            self.drag_start_pos = None
            self.dragged_text = None


class IssueReportingApp(tk.Tk):
    def __init__(self):
        super().__init__()

        # Initialize UI elements
        self.init_ui()

        # List to store images
        self.images = []

    def init_ui(self):


        self.content_frame = tk.Frame(self)
        self.content_frame.pack(side='left', fill='both', expand=True)

        # Create input fields for the required information
        ttk.Label(self.content_frame, text="Req-Number").grid(row=0, column=0, sticky='w', pady=5)
        self.req_number = ttk.Entry(self.content_frame)
        self.req_number.grid(row=0, column=1, pady=5)

        ttk.Label(self.content_frame, text="Company").grid(row=1, column=0, sticky='w', pady=5)
        self.company = ttk.Entry(self.content_frame)
        self.company.grid(row=1, column=1, pady=5)

        ttk.Label(self.content_frame, text="Country").grid(row=2, column=0, sticky='w', pady=5)
        self.country = ttk.Combobox(self.content_frame, values=['USA', 'UK', 'Germany', 'France']) # Add more countries as needed
        self.country.grid(row=2, column=1, pady=5)

        ttk.Label(self.content_frame, text="Account name").grid(row=3, column=0, sticky='w', pady=5)
        self.account_name = ttk.Entry(self.content_frame)
        self.account_name.grid(row=3, column=1, pady=5)

        ttk.Label(self.content_frame, text="Impacted user's ID(s)").grid(row=4, column=0, sticky='w', pady=5)
        self.user_ids = ttk.Entry(self.content_frame)
        self.user_ids.grid(row=4, column=1, pady=5)

        ttk.Label(self.content_frame, text="Environment, Application, Sub Product, Dataset").grid(row=5, column=0, sticky='w', pady=5)
        self.dataset = ttk.Entry(self.content_frame)
        self.dataset.grid(row=5, column=1, pady=5)

        ttk.Label(self.content_frame, text="Report Details").grid(row=6, column=0, sticky='w', pady=5)
        self.report_details = CustomScrolledText(self.content_frame, width=40, height=5)
        self.report_details.grid(row=6, column=1, pady=5)

        ttk.Label(self.content_frame, text="Is the issue replicable?").grid(row=7, column=0, sticky='w', pady=5)
        self.replicable = ttk.Combobox(self.content_frame, values=['YES', 'NO'])
        self.replicable.grid(row=7, column=1, pady=5)

        ttk.Label(self.content_frame, text="Steps/Troubleshooting").grid(row=8, column=0, sticky='w', pady=5)
        self.steps = CustomScrolledText(self.content_frame, width=40, height=5)
        self.steps.grid(row=8, column=1, pady=5)

        ttk.Label(self.content_frame, text="Time and timezone of error").grid(row=9, column=0, sticky='w', pady=5)
        self.error_time = ttk.Entry(self.content_frame)
        self.error_time.grid(row=9,column=1, pady=5)

        ttk.Label(self.content_frame, text="Describe the issue").grid(row=10, column=0, sticky='w', pady=5)
        self.issue_description = DragDropText(self.content_frame, width=40, height=5)
        self.issue_description.grid(row=10, column=1, pady=5)

        # Image display area
        self.image_label = tk.Label(self.content_frame) 
        self.image_label.grid(row=11, column=0, columnspan=2)

        # Bind Ctrl+V for pasting images
        self.bind('<Control-v>', self.paste_screenshot)

        # Create buttons for generating output
        self.generate_button = ttk.Button(self.content_frame, text="Generate Word Document", command=self.generate_word_document)
        self.generate_button.grid(row=12, column=0, columnspan=2, pady=5)

        self.copy_output_button = ttk.Button(self.content_frame, text="Copy Output", command=self.generate_copy_ready_text)
        self.copy_output_button.grid(row=13, column=0, columnspan=2, pady=5)

        # Create a button for exiting the application
        self.exit_button = ttk.Button(self.content_frame, text="Exit", command=self.exit_application)
        self.exit_button.grid(row=14,column=0, columnspan=2, pady=5)

    def paste_screenshot(self, event):
        try:
            clipboard_content = clipboard.paste()
            # Replace forward slashes with backslashes in the file path
            clipboard_content = clipboard_content.replace('/', '\\')
            if os.path.isfile(clipboard_content):
                # Open the image file and display it
                image = Image.open(clipboard_content)
                # Save the image and add it to the list
                new_filename = os.path.abspath(f"image_{len(self.images)}.png")
                image.save(new_filename)
                self.images.append(new_filename)

                # Copy the image file to the application's folder
                shutil.copy(clipboard_content, new_filename)

                # Insert placeholder into active text field
                active_widget = self.focus_get()
                if isinstance(active_widget, CustomScrolledText):
                    active_widget.insert(tk.END, f"\n{{{new_filename}}}\n") # Always wrap filename in curly braces
            else:
                # Handle pasted image data
                image = ImageGrab.grabclipboard()
                if isinstance(image, Image.Image): # Make sure it's an image
                    # Save the image and add it to the list
                    filename = os.path.abspath(f"image_{len(self.images)}.png")
                    image.save(filename)
                    self.images.append(filename)

                    # Insert placeholder into active text field
                    active_widget = self.focus_get()
                    if isinstance(active_widget, CustomScrolledText):
                        active_widget.insert(tk.END, f"\n{{{filename}}}\n") # Always wrap filename in curly braces
        except Exception as e:
            print(e) # This will print the actual error message
            messagebox.showerror('Error', 'Could not paste the image!')


    def add_image(self):
        filename = filedialog.askopenfilename(initialdir="/", title="Select Image", filetypes=(("jpeg files", "*.jpg"), ("png files", "*.png")))
        if filename:
            image = Image.open(filename)
            image.thumbnail((200, 200))
            photo = ImageTk.PhotoImage(image)

            self.images.append(filename)
            label = ImageLabel(self.content_frame, image_filename=filename, image=photo)
            label.image = photo
            label.pack()

    def generate_word_document(self):
        doc = docx.Document()

        # Add content to the document
        doc.add_paragraph(f"Req-Number: {self.req_number.get()}")
        doc.add_paragraph(f"Company: {self.company.get()}")
        doc.add_paragraph(f"Country: {self.country.get()}")
        doc.add_paragraph(f"Account name: {self.account_name.get()}")
        doc.add_paragraph(f"Impacted user's ID(s): {self.user_ids.get()}")
        doc.add_paragraph(f"Environment, Application, Sub Product, Dataset: {self.dataset.get()}")
        
        # Parse text from report_details text field and insert images at appropriate locations
        text = self.report_details.get('1.0', 'end')
        lines = text.split('\n')
        for line in lines:
            # Check if line is a file path wrapped in curly braces
            if line.startswith('{') and line.endswith('}'):
                image_filename = line[1:-1]  # Remove the curly braces
            else:
                image_filename = line  # Use the line as is

            # Replace forward slashes with backslashes in the file path
            image_filename = image_filename.replace('/', '\\')

            # Create a new paragraph and a new run
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()

            if os.path.isfile(image_filename):
                # Add the image to the run
                run.add_picture(image_filename, width=Inches(6.0))  # Adjust the width as needed
            else:
                # Add the text to the run
                run.add_text(line)
        

        #ADD THIS TO ANOTHER PLACE!!!!!!
        # Save the Word document to a temporary file
        temp_filename = 'temp.docx'
        doc.save(temp_filename)

        # Open the Word document and copy its content to the clipboard
        word = win32com.client.Dispatch('Word.Application')
        doc = word.Documents.Open(os.path.abspath(temp_filename))
        doc.Content.Copy()
        doc.Close()

        # Delete the temporary file
        os.remove(temp_filename)

        messagebox.showinfo('Info', 'Content copied to clipboard!')
        


        # Save the Word document
        try:
            doc.save(self.req_number.get() + '.docx')
            messagebox.showinfo('Info', 'Word document created successfully!')
        except Exception as e:
            messagebox.showerror('Error', 'Could not save the Word document!')





    def generate_copy_ready_text(self):
        # Prepare text using input fields
        output = f"Req-Number: {self.req_number.get()}\n"
        output += f"Company: {self.company.get()}\n"
        output += f"Country: {self.country.get()}\n"
        output += f"Account name: {self.account_name.get()}\n"
        output += f"Impacted user's ID(s): {self.user_ids.get()}\n"
        output += f"Environment, Application, Sub Product, Dataset: {self.dataset.get()}\n"
        output += f"Report Details: {self.report_details.get('1.0', 'end')}\n"
        output += f"Is the issue replicable?: {self.replicable.get()}\n"
        output += f"Steps/Troubleshooting: {self.steps.get('1.0', 'end')}\n"
        output += f"Time and timezone of error: {self.error_time.get()}\n"
        output += f"Describe the issue: {self.issue_description.get('1.0', 'end')}\n"


        # Copy text to clipboard
        clipboard.copy(output)
        messagebox.showinfo('Info', 'Text copied to clipboard!')

    def exit_application(self):
        msg_box = messagebox.askquestion('Exit Application', 'Are you sure you want to exit the application?', icon='warning')
        if msg_box == 'yes':
            self.destroy()
        else:
            messagebox.showinfo('Return', 'You will now return to the application screen')

if __name__ == '__main__':
    app = IssueReportingApp()
    app.mainloop()


